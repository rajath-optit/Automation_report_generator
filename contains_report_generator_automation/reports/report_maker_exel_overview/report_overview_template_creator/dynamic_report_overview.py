from docx import Document
from datetime import datetime

def create_report_overview_template(output_file):
    # Get client name from user input
    client_name = input("Enter the client name: ").strip()
    report_name = f"AWS {client_name} Report"
    
    # Get the current date
    today_date = datetime.now().strftime('%A, %dth %B %Y')

    # Initialize the document
    doc = Document()

    # Add the report title
    doc.add_heading(report_name, level=1)

    # Add client and report details
    doc.add_paragraph(f"- Client: {client_name}")
    doc.add_paragraph(f"- Report Name: {report_name}")
    doc.add_paragraph(f"- Report Date: {today_date}")
    doc.add_paragraph("- Report Version: Version 1.0")

    # Add a section break
    doc.add_paragraph("\n")

    # Add index section
    doc.add_heading("Index:", level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Title"
    table.cell(0, 1).text = "Page No"
    table.cell(1, 0).text = "Overview"
    table.cell(1, 1).text = "03"
    table.cell(2, 0).text = "LINK OF DETAILED REPORT"
    table.cell(2, 1).text = "04"
    table.cell(3, 0).text = "AWS Compliance Control Summary"
    table.cell(3, 1).text = "06"
    table.cell(4, 0).text = "Open Issues Chart"
    table.cell(4, 1).text = "47"
    table.cell(5, 0).text = "Top services severity"
    table.cell(5, 1).text = "48"
    doc.add_paragraph("(Page numbers to be filled by the user)")

    # Add the Overview section
    doc.add_heading("Overview:", level=2)
    doc.add_paragraph(
        "The AWS Compliance module evaluates your AWS account against frameworks such as CIS, GDPR, PCI, and others. "
        "This report provides insights through detailed benchmarks and controls, addressing key questions:"
    )
    doc.add_paragraph("1. Resource Inventory: Quantifies resources by type.")
    doc.add_paragraph("2. Counts by Accounts/Regions: Breaks down resources by accounts and locations.")
    doc.add_paragraph("3. Configuration Percentages: Tracks configurations (e.g., encryption enabled).")
    doc.add_paragraph("4. Resource Age: Identifies age for lifecycle management.")
    doc.add_paragraph(
        "For further insights and detailed resource analysis, refer to the Excel sheet provided."
    )

    # Add a placeholder for a link to the detailed report
    doc.add_heading("LINK OF DETAILED REPORT:", level=2)
    doc.add_paragraph("[Add the link here]")

    # Add the Key Components table
    doc.add_heading("Key Components:", level=2)
    components_table = doc.add_table(rows=9, cols=2)
    components_table.style = "Table Grid"
    components_table.cell(0, 0).text = "S.No"
    components_table.cell(0, 1).text = "Category"
    components_table.cell(1, 0).text = "1"
    components_table.cell(1, 1).text = "Sheet Name"
    components_table.cell(2, 0).text = ""
    components_table.cell(2, 1).text = "Compute Services"
    components_table.cell(3, 0).text = ""
    components_table.cell(3, 1).text = "Compute"
    components_table.cell(4, 0).text = ""
    components_table.cell(4, 1).text = "Storage Services"
    components_table.cell(5, 0).text = ""
    components_table.cell(5, 1).text = "Storage"
    components_table.cell(6, 0).text = ""
    components_table.cell(6, 1).text = "Network Services"
    components_table.cell(7, 0).text = ""
    components_table.cell(7, 1).text = "Database Services"
    components_table.cell(8, 0).text = ""
    components_table.cell(8, 1).text = "Security Services"

    # Add the AWS Compliance Control Summary
    doc.add_heading("AWS Compliance Control Summary", level=2)
    doc.add_paragraph(
        "This table provides an overview of critical compliance controls for Auto Scaling, EC2, S3, and more. "
        "It includes descriptions, open issues, and priority levels to ensure security and operational efficiency. "
        "Refer to the attached Excel sheet for details."
    )

    # Add placeholders for user inputs (e.g., tables and summary details)
    doc.add_paragraph("[Insert table with compliance controls here]")
    doc.add_paragraph(
        "The Total Count of Titles is 1,554 for priority level 1, 3,328 for priority level 2, "
        "and 1,847 for priority level 3, resulting in a Grand Total of 6,729."
    )

    # Add Synopsis section
    doc.add_heading("SYNOPSIS:", level=2)
    doc.add_paragraph("[Add synopsis here]")

    # Save the document
    doc.save(output_file)
    print(f"Report overview template saved as {output_file}")


if __name__ == "__main__":
    output_file = "AWS_Report_Overview_Template.docx"
    create_report_overview_template(output_file)
